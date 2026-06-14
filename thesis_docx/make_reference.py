"""
make_reference.py — build a styled reference.docx encoding the thesis formatting
requirements (Innopolis MS standard). Pandoc uses this file as the style source
when converting the Markdown sources to the final thesis.docx.

Requirements encoded:
  - Times New Roman, 14 pt body
  - line spacing 1.5
  - margins: left 25 mm, right/top/bottom 20 mm
  - first-line indent 1.25 cm, justified body
  - each chapter (Heading 1) starts on a new page
  - continuous page numbering in the footer (centered)
  - figure caption left-aligned, table caption centered
"""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(style, name="Times New Roman"):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


def _ensure_style(doc, name, base="Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles[base]
        return st


doc = Document()

# ── Normal (body) ────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
_set_font(normal)
normal.font.size = Pt(14)
normal.font.color.rgb = RGBColor(0, 0, 0)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.first_line_indent = Cm(1.25)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ── Headings (TNR, black, bold) ──────────────────────────────────────────────
for lvl, size in ((1, 16), (2, 15), (3, 14)):
    h = doc.styles[f"Heading {lvl}"]
    _set_font(h)
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    hpf = h.paragraph_format
    hpf.first_line_indent = Cm(0)
    hpf.space_before = Pt(12)
    hpf.space_after = Pt(6)
    hpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
# Each chapter on a new page
doc.styles["Heading 1"].paragraph_format.page_break_before = True

# ── Title / Author ───────────────────────────────────────────────────────────
for nm in ("Title", "Subtitle"):
    try:
        _set_font(doc.styles[nm])
    except KeyError:
        pass

# ── Captions: figure left, table centered ────────────────────────────────────
img_cap = _ensure_style(doc, "Image Caption")
_set_font(img_cap); img_cap.font.size = Pt(12)
img_cap.paragraph_format.first_line_indent = Cm(0)
img_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

tbl_cap = _ensure_style(doc, "Table Caption")
_set_font(tbl_cap); tbl_cap.font.size = Pt(12)
tbl_cap.paragraph_format.first_line_indent = Cm(0)
tbl_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Generic pandoc "Caption" style (fallback)
cap = _ensure_style(doc, "Caption")
_set_font(cap); cap.font.size = Pt(12)
cap.paragraph_format.first_line_indent = Cm(0)

# ── Source code (monospace, no indent) ───────────────────────────────────────
src = _ensure_style(doc, "Source Code")
src.font.name = "Consolas"; src.font.size = Pt(10)
src.paragraph_format.first_line_indent = Cm(0)

# ── Page setup + footer page number ──────────────────────────────────────────
sec = doc.sections[0]
sec.left_margin = Mm(25)
sec.right_margin = Mm(20)
sec.top_margin = Mm(20)
sec.bottom_margin = Mm(20)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.first_line_indent = Cm(0)
_add_page_number_field(fp)
_set_font(doc.styles["Footer"])

import os
out = os.path.join(os.path.dirname(__file__), "reference.docx")
doc.save(out)
print("wrote", out)
