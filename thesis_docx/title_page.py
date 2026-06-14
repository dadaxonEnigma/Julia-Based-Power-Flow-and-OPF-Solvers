"""
title_page.py — inject the Innopolis title page into the built thesis.docx and
configure page numbering so the title page is counted but unnumbered, with
continuous numbering from page 2.

Called by build.py after pandoc conversion. Edit TITLE_PAGE_LINES to change the
title-page content.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

# (text, font_size_pt, bold, space_before_pt)  — rendered centered, no indent.
TITLE_PAGE_LINES = [
    ("Innopolis University", 16, True, 24),
    ("Master's Programme in Computer Science", 14, False, 6),
    ("AI-Assisted Simulation and Optimization of Power Market Projects", 20, True, 90),
    ("Master's Thesis", 14, False, 36),
    ("Author: Dadakhon Turgunboev", 14, False, 60),
    ("Supervisor: Leonard Johard", 14, False, 6),
    ("Innopolis, 2026", 14, False, 90),
]


def _make_paragraph(doc, text, size, bold, space_before):
    p = doc.add_paragraph()                       # appended at end (moved later)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def add_title_page(path: str):
    doc = Document(path)
    body = doc.element.body

    # Build the title-page paragraphs (temporarily at the end of the document).
    pars = [_make_paragraph(doc, t, s, b, sb) for (t, s, b, sb) in TITLE_PAGE_LINES]
    # Trailing paragraph carrying a page break so the TOC starts on a new page.
    brk = doc.add_paragraph()
    brk.paragraph_format.first_line_indent = Cm(0)
    brk.add_run().add_break(WD_BREAK.PAGE)
    pars.append(brk)

    # Move them to the very top, before the pandoc-generated TOC, preserving order.
    for p in reversed(pars):
        body.insert(0, p._p)

    # Title page counted but unnumbered: a different (empty) first-page footer,
    # while the normal footer keeps its PAGE field for pages 2+.
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    fpf = sec.first_page_footer
    if fpf.paragraphs:
        fpf.paragraphs[0].text = ""             # ensure no number on the title page

    doc.save(path)
